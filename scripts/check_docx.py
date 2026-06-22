"""检查 docx 文件内容"""
from docx import Document
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

docx_path = r"D:\我的资料\书籍\大师兄.docx"

try:
    doc = Document(docx_path)
    print(f"=== 大师兄.docx ===")
    print(f"总段落数: {len(doc.paragraphs)}")
    print()

    # 打印前 50 个非空段落
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] {para.text[:200]}")
            count += 1
            if count >= 50:
                break

    print()
    print(f"前 50 个非空段落已显示")

except Exception as e:
    print(f"错误: {e}")
    print("尝试安装 python-docx: pip install python-docx")
